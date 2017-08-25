# -*- coding: utf-8 -*-
"""
Created on Fri Mar 18 14:21:12 2016

@author: k
"""
def combineFilesDm():
    folder = "F:\\Insects\\Anopheles_gambiae\\Interpro\\interpro\\"
    import glob
    filenames = glob.glob(folder+"*.tsv")
    fout = open(folder+"AgMCOT_Interpro.tsv","w")
    for file in filenames:
        fout.write(open(file).read())
    fout.close()

def getSPSeqs():
    from Bio import SeqIO
    dcFa = SeqIO.to_dict(SeqIO.parse(open("F:\Insects\Drosophila_melanogaster\\dmel-all-translation-r6.09.fasta"),"fasta"))
    mylist = open(r"F:\Insects\Drosophila_melanogaster\SerineProteases\list.txt").read().split()
    fout = open(r"F:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aa.txt","w")
    for ele in mylist:
        if ele in dcFa:
            fout.write(">"+dcFa[ele].id +" "+dcFa[ele].description.split()[4][5:-1]+" "+dcFa[ele].description.split()[5]\
            +" " +dcFa[ele].description.split()[8]+"\n"+str(dcFa[ele].seq)+"\n")
    fout.close()

def getSPSeqs2():
    from Bio import SeqIO
    dcFa = SeqIO.to_dict(SeqIO.parse(open(r"F:\Insects\Anopheles_gambiae\Anopheles-gambiae-PEST_PEPTIDES_AgamP4.3.fa"),"fasta"))
    mylist = open(r"F:\Insects\Anopheles_gambiae\SPSPH\list.txt").read().split()
    fout = open(r"F:\Insects\Anopheles_gambiae\SPSPH\SPSPH_aa.txt","w")
    for ele in mylist:
        if ele in dcFa:
            fout.write(">"+dcFa[ele].description+"\n"+str(dcFa[ele].seq)+"\n")
    fout.close()

def getSPSeqs3():
    from Bio import SeqIO
    dcFa = SeqIO.to_dict(SeqIO.parse(open(r"F:\Insects\Anopheles_gambiae\mcot\20160316AgMcotprotein.fa"),"fasta"))
    mylist = open(r"F:\Insects\Anopheles_gambiae\SPSPH\MCOTlist.txt").read().split()
    fout = open(r"F:\Insects\Anopheles_gambiae\SPSPH\MCOTSPSPH_aa.txt","w")
    for ele in mylist:
        if ele in dcFa:
            fout.write(">"+dcFa[ele].description+"\n"+str(dcFa[ele].seq)+"\n")
    fout.close()


def testDocx():
    mylist = open(r"F:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aa.txt").readlines()
    from docx import Document
    from docx.shared import RGBColor
    document = Document()
    for line in mylist:
        p = document.add_paragraph()
        if line[0] == ">":
            run = p.add_run(line[:-1])
            font = run.font
            font.color.rgb = RGBColor(0xFF,0x00,0x00)
            font.name = "Courier New"
        else:
            run = p.add_run(line[:-1])
            font = run.font
            font.color.rgb = None
            font.name = "Courier New"
    document.save("test.docx")

def tsvSelect():
    folder = "F:\\Insects\\Drosophila_melanogaster\\interpro\\"
    mylist = open(r"F:\Insects\Drosophila_melanogaster\SerineProteases\list.txt").read().split()
    mytsv = open(folder+"DmeInterpro.tsv").readlines()
    fout = open(r"F:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_tsvDomain.txt","w")
    for line in mytsv:
        if line.split()[0] in mylist:
            fout.write(line)
    fout.close()

def tsvSelect2():
    folder = "F:\\Insects\\Anopheles_gambiae\\Interpro\\"
    mylist = open(r"F:\Insects\Anopheles_gambiae\SPSPH\list.txt").read().split()
    mytsv = open(folder+"AgInterpro.tsv").readlines()
    fout = open(r"F:\Insects\Anopheles_gambiae\SPSPH\SPSPH_tsvDomain.txt","w")
    for line in mytsv:
        if line.split()[0] in mylist:
            fout.write(line)
    fout.close()

def tsvSelect3():
    folder = "F:\\Insects\\Anopheles_gambiae\\Interpro\\"
    mylist = open(r"F:\Insects\Anopheles_gambiae\SPSPH\MCOTlist.txt").read().split()
    mytsv = open(folder+"AgMCOT_Interpro.tsv").readlines()
    fout = open(r"F:\Insects\Anopheles_gambiae\SPSPH\MCOTSPSPH_tsvDomain.txt","w")
    for line in mytsv:
        if line.split()[0] in mylist:
            fout.write(line)
    fout.close()




def save2DocxWithSmartDomain(fileSPfa=None,filelsDomainSP=None):
    """
    save the text file of fasta format to word, with smart domains labeled.
    """
    if fileSPfa is None:
        fileSPfa = r"F:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aa.txt"
    if filelsDomainSP is None:
        filelsDomainSP = r"F:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_tsvDomain.txt"
    from Bio import SeqIO
    dcFa = SeqIO.to_dict(SeqIO.parse(open(fileSPfa),"fasta"))
    lsFa = list(SeqIO.parse(open(fileSPfa),"fasta"))
    lsdomain = open(filelsDomainSP).readlines()
    dcDomain ={}
    for ele in lsdomain:#store domain information in a dictionary
        ele_detail = ele.split("\t")
        if ele_detail[0] not in dcDomain:
            dcDomain[ele_detail[0]] =[]
        dcDomain[ele_detail[0]].append(ele_detail)
    dcSignalP = {}#store signalP location
    for key in dcDomain:
        for ele in dcDomain[key]:
            if "SignalP_EUK" == ele[3]:
                dcSignalP[ele[0]] = ["SignalP",int(ele[6])-1,int(ele[7]),0]
    dcSmart = {}#store SMART domain location
    for key in dcDomain:
        for ele in dcDomain[key]:
            if "SMART" == ele[3]:
                if key not in dcSmart:
                    dcSmart[key] = []
                dcSmart[key].append([ele[5],int(ele[6])-1,int(ele[7]),float(ele[8])])
    dcSlices ={}
    for key in dcFa:#combin dcSignalP and dcSmart
        dcSlices[key] = []
        if key in dcSignalP:
            dcSlices[key].append(dcSignalP[key])
        if key in dcSmart:
            for ele in dcSmart[key]:
                dcSlices[key].append(ele)
    for key in dcSlices:#sort by starting site
        dcSlices[key] = sorted(dcSlices[key],key = lambda x: x[1])
    
    dcSP = {}#store SP domain
    for key in dcDomain:
        for ele in dcDomain[key]:
            if "SM00020" == ele[4]:
                dcSP[key] = [ele[5],int(ele[6])-1,int(ele[7]),float(ele[8])]
    
    dcDocSlices ={}
    for key in lsFa:
        key = key.id
        dcDocSlices[key] = []
        if key in dcSignalP:
            if key in dcSP:
                if dcSignalP[key][2]>dcSP[key][1]:
                    dcSP[key][1] =dcSignalP[key][2]
                dcDocSlices[key] = [dcSignalP[key],["",dcSignalP[key][2], dcSP[key][1]],dcSP[key],["",dcSP[key][2],len(dcFa[key])]]
            else:
                dcDocSlices[key] = [dcSignalP[key],["",dcSignalP[key][2],len(dcFa[key])]]
        else:
            if key in dcSP:
                dcDocSlices[key] = [["",0, dcSP[key][1]],dcSP[key],["",dcSP[key][2],len(dcFa[key])]]
            else:
                dcDocSlices[key] = [["",0,len(dcFa[key])]]
    
                    
    from docx import Document
    from docx.shared import RGBColor
    document = Document()
    for key in lsFa:
        key = key.id
        p = document.add_paragraph()#fasta head
        run = p.add_run("\n\n\n>"+dcFa[key].description)
        font = run.font
        font.color.rgb = RGBColor(0x00,0x00,0x00)
        font.name = "Courier New"
        
        p = document.add_paragraph()#domains and signalP basic
        for ele in dcSlices[key]:
            run = p.add_run(ele[0]+"\t")
            font = run.font
            font.color.rgb = RGBColor(0x00,0x00,0x00)
            font.name = "Courier New"
        
        seq = str(dcFa[key].seq)
        
        p = document.add_paragraph()#domain detail
        for ele in dcSlices[key]:
            if ele not in dcDocSlices[key]:
                run = p.add_run(ele[0] +"\t"+seq[ele[1]:ele[2]]+"\n")
                font = run.font
                font.color.rgb = RGBColor(0x00,0x00,0xFF)
                font.name = "Courier New"
        
        p = document.add_paragraph()#write seq
        for ele in dcDocSlices[key]:
            if "SignalP" in ele[0]:
                run = p.add_run(seq[ele[1]:ele[2]])
                font = run.font
                font.color.rgb = RGBColor(0x00,0xFF,0x00)
                font.name = "Courier New"
            elif "protease" in ele[0]:
                run = p.add_run(seq[ele[1]:ele[2]])
                font = run.font
                font.color.rgb = RGBColor(0xFF,0x00,0x00)
                font.name = "Courier New"
            else:
                run = p.add_run(seq[ele[1]:ele[2]])
                font = run.font
                font.color.rgb = RGBColor(0x00,0x00,0x00)
                font.name = "Courier New"
#        break
    document.save("domainSP.docx")

def run():
    fileSPfa = r"F:\Insects\Anopheles_gambiae\SPSPH\SPSPH_aa.txt"
    filelsDomainSP = r"F:\Insects\Anopheles_gambiae\SPSPH\SPSPH_tsvDomain.txt"
    save2DocxWithSmartDomain(fileSPfa,filelsDomainSP)
            
    fileSPfa = r"F:\Insects\Anopheles_gambiae\SPSPH\MCOTSPSPH_aa.txt"
    filelsDomainSP = r"F:\Insects\Anopheles_gambiae\SPSPH\MCOTSPSPH_tsvDomain.txt"
    save2DocxWithSmartDomain(fileSPfa,filelsDomainSP)
    
    
    fileSPfa = r"F:\Insects\Anopheles_gambiae\SPSPH\SPSPH_combinedAA.txt"
    filelsDomainSP = r"F:\Insects\Anopheles_gambiae\SPSPH\SPSPH_combinedDomain.txt"
    save2DocxWithSmartDomain(fileSPfa,filelsDomainSP)

def compareOGS():
    '''
    compare difference of Dm6.09 vs Dm6.14
    '''
    f_43 = r"D:\Insects\Drosophila_melanogaster\dmel-all-translation-r6.09.fasta"
    f_44 = r"D:\Insects\Drosophila_melanogaster\dmel-all-translation-r6.14.fasta"
    import largeFastaFile
    dc43 = largeFastaFile.open_fasta_to_dic(f_43)
    dc44 = largeFastaFile.open_fasta_to_dic(f_44)
    
    
    ls_both = []
    ls_43only = []
    ls_44only = []
    keys = set()
    keys.update(dc43.keys())
    keys.update(dc44.keys())
    for key in keys:
        if key in dc43:
            if key in dc44:
                ls_both.append(key)
            else:
                ls_43only.append(key)
        else:
            if key in dc44:
                ls_44only.append(key)
    
    ls_both_identical = []
    ls_both_nonIdentical = []
    for key in ls_both:
        if str(dc43[key].seq) == str(dc44[key].seq):
            ls_both_identical.append(key)
        else:
            ls_both_nonIdentical.append(key)
    
    print(f_43, 'number of Pr:', len(dc43))
    print(f_44, 'number of Pr:', len(dc44))
    
    print(f_43, 'only Pr:', len(ls_43only))
    print(f_44, 'only Pr:', len(ls_44only))
    
    print('both, identical Pr:', len(ls_both_identical))
    print('both, non-identical Pr.', len(ls_both_nonIdentical))
    
    #check if SPs in ls_43only, 
    ls_SP = open(r"D:\Insects\Drosophila_melanogaster\SerineProteases\list.txt").read().split()
    print(len([ele for ele in ls_SP if ele in ls_43only]))
    #6.09 only seqs, none of them is SP
    #check 6.14 only seqs
    fout = open('seq.txt','w')
    for ele in ls_44only:
        fout.write('>'+dc44[ele].id +'\n'+str(dc44[ele].seq) +'\n')
    fout.close()
    #blast result shows that none sequence ins ls_44only is serine protease. No change to make over the previous result.

def compareSPSPHwithTrintiy20170309():
    '''
    compare the SPSPH sequences from OGS with Trinity.
    '''
    import largeFastaFile
    lsSPs = largeFastaFile.open_fasta_to_list(r"D:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aa.txt")
    lsTrinity = largeFastaFile.open_fasta_to_list(r"D:\Insects\Drosophila_melanogaster\2017GeneModeling\DmTrinity.fa.transdecoder.pep")
    print('SP number is ', len(lsSPs))
    print('Pr in Trinity is ', len(lsTrinity))
    
    
    stSeq = {str(ele.seq).replace('*','') for ele in lsTrinity}
    print('unique seqs in Trinity is ', len(stSeq))
    
    lsNonIdentical = [ele for ele in lsSPs if str(ele.seq) not in stSeq]
    print('None identitial SPs is ', len(lsNonIdentical))
    
    
    #use the blast result
    #read in lsTrinityUnique
    dcTrinityU = largeFastaFile.open_fasta_to_dic(r"D:\Insects\Drosophila_melanogaster\2017GeneModeling\DmTrinityUniqueNonWithin20170309.txt")
    #read in blast result in table format. only keep top line with subject id in dcTrinityU
    dcSP2T = {}
    for line in open("D:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aaTab.txt"):
        eles = line.split()
        _sp = eles[0]
        _t = eles[1]
        if _t in dcTrinityU and _sp not in dcSP2T:
            dcSP2T[_sp] = _t
    print('SP with hit, number is ', len(dcSP2T))
    
    dcSP = largeFastaFile.open_fasta_to_dic(r"D:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aa.txt")
    for _spid,_tid in dcSP2T.items():
        _sp = dcSP[_spid]
        _t = dcTrinityU[_tid]
        _splen = len(sp.seq)
        _tlen = len(str(_t.seq).replace('*',''))
    
    #blast with edited scoreing matrix
    #calculate matched length with the function used in McuNovo
    #336 SPs, 20 have no hit. 175 are identical to Trinity. 
    #Of the 141 non-identical ones, 80 is SPs longer, 61 is Trinity longer
    #do a one 2 one alignement for these sequences
    ls_SPTpair = []
    for ele in open(r"D:\Insects\Drosophila_melanogaster\SerineProteases\20170313SPSPH_withTrinityPair_non_identical_IDs.txt"):
        ls_SPTpair.append((ele.split()[0],ele.split()[2]))
    MUSCLE = r"C:\Users\k\AppData\Roaming\MEGA7_7160630-x86_64\Private\MUSCLE\muscleWin64.exe"
    import subprocess
    file = r"D:\P\3Language\Xiaolong\python\seq.txt"
    for ele in ls_SPTpair:
        fo = open(file,'w')
        _sp = dcSP[ele[0]]
        _t = dcTrinityU[ele[1]]
        fo.write('>'+_sp.id+'\n'+str(_sp.seq)+'\n')
        fo.write('>'+_t.id+'\n'+str(_t.seq)+'\n')
        fo.close()
        subprocess.run('echo %s >>D:\\P\\3Language\\Xiaolong\\python\\algin.txt'%_t.description,shell=True)
        subprocess.run(MUSCLE + ' -clw -in '+file+ ' >>D:\\P\\3Language\\Xiaolong\\python\\algin.txt',shell=True)
        

def check260muscleAlignement20170320():
    '''
    check the alignments
    '''
    #get protein2gene file
    fogspep = r"D:\Insects\Drosophila_melanogaster\dmel-all-translation-r6.14.fasta"
    import largeFastaFile
    ls_ogspep = largeFastaFile.open_fasta_to_list(fogspep)
    fout = open(r"D:\Insects\Drosophila_melanogaster\dmel-all-translation-r6.14ProteinInfo.txt",'w')
    for ele in ls_ogspep:
        _des = ele.description.split()
        _gene = _des[5][7:18]
        _chr = _des[2].split('=')[1].split(':')[0]
        if 'complement' in _des[2]:
            _strand = '-'
        else:
            _strand = '+'
        if 'join' in _des[2]:
            _exons = _des[2].split('join')[1][1:].split(')')[0]
            _start = _exons.split(',')[0].split('..')[0]
            _end = _exons.split(',')[-1].split('..')[-1]
            _exonscount = str(len(_exons.split(',')))
        else:
            _exonscount = '1'
            if 'complement' not in _des[2]:
                _start = _des[2].split(':')[1].split('..')[0]
                _end = _des[2].split(':')[1].split('..')[1][:-1]
            else:
                _start = _des[2].split(':')[1].split('..')[0][11:]
                _end = _des[2].split(':')[1].split('..')[1][:-2]
        fout.write('\t'.join([ele.id, _gene, _chr, _exonscount, _strand, _start, _end]) +'\n')
    fout.close()
    
    #count location with no gaps, alignemtn
    falign = r"D:\Insects\Drosophila_melanogaster\SerineProteases\20170320Tree\20170321DmSPs249.fas"
    ls_align = largeFastaFile.open_fasta_to_list(falign)
    ls_align = [str(ele.seq) for ele in ls_align]
    count = 0
    for n in range(len(ls_align[0])):
        _position = [ele[n] for ele in ls_align]
        if '-' not in _position:
            count += 1
    print(count)
    
    #remove alignment columns with more than 240 gap symbols
    falign = r"D:\Insects\Drosophila_melanogaster\SerineProteases\20170320Tree\20170321DmSPs249.fas"
    ls_align = largeFastaFile.open_fasta_to_list(falign)
    ls_seq = [str(ele.seq) for ele in ls_align]
    to_remove = []
    for n in range(len(ls_align[0])):
        _position = [ele[n] for ele in ls_align]
        if _position.count('-') >240:
            to_remove.append(n)
    
    to_remove = set(to_remove)
    fout = open(falign+'240.fas','w')
    for n in range(len(ls_align)):
        fout.write('>'+ls_align[n].id+'\n')
        for m in range(len(ls_align[0])):
            if m not in to_remove:
                fout.write(ls_seq[n][m])
        fout.write('\n')
    fout.close()
    
    #get chromosome length
    f_chr = r"D:\Insects\Drosophila_melanogaster\dmel-all-chromosome-r6.14.fasta"
    import largeFastaFile
    ls_chr = largeFastaFile.open_fasta_to_list(f_chr)
    fout = open('list.txt','w')
    for ele in ls_chr:
        fout.write(ele.id+'\t'+str(len(ele.seq))+'\n')
    fout.close()
    
    #modify alignement to include more information
    
    
def rsemFromSRR20170322():
    '''
    generate RSEM qsub file based on SRR numbers
    '''
    ls = open('list.txt').readlines()
    print(len(ls))
    def srr2wgetEBI(srr,paired = True):
        '''
        given a srr number, return a line of srr download script
        '''
        if len(srr) == 9:
            if not paired:
                return 'wget ftp://ftp.sra.ebi.ac.uk/vol1/fastq/{srr6}/{srr}/{srr}.fastq.gz -q \n'.format(srr=srr, srr6=srr[:6])
            else:
                return 'wget ftp://ftp.sra.ebi.ac.uk/vol1/fastq/{srr6}/{srr}/{srr}_1.fastq.gz -q &\nwget ftp://ftp.sra.ebi.ac.uk/vol1/fastq/{srr6}/{srr}/{srr}_2.fastq.gz -q \nwait $(jobs -p)\n'.format(srr=srr, srr6=srr[:6])
        else:
            if not paired:
                return 'wget ftp://ftp.sra.ebi.ac.uk/vol1/fastq/{srr6}/00{srr1}/{srr}/{srr}.fastq.gz -q \n'.format(srr=srr, srr6=srr[:6], srr1 = srr[-1])
            else:
                return 'wget ftp://ftp.sra.ebi.ac.uk/vol1/fastq/{srr6}/00{srr1}/{srr}/{srr}_1.fastq.gz -q &\nwget ftp://ftp.sra.ebi.ac.uk/vol1/fastq/{srr6}/00{srr1}/{srr}/{srr}_2.fastq.gz -q \nwait $(jobs -p)\n'.format(srr=srr, srr6=srr[:6], srr1 = srr[-1])
                
    for line in ls[1:]:
        _es = line.split('\t')
        _srr = _es[1]
        _p = _es[3]
        fout = open('2017DmRSEM_%s.txt'%(_srr),'w')
        fout.write('''
#!/bin/bash
#PBS -q batch
#PBS -N {srr}
#PBS -l nodes=1:ppn=12
#PBS -l walltime=120:00:00
#PBS -m abe -M atps@outlook.com
#PBS -j oe
cd /panfs/panfs.cluster/scratch/ks2073/RSEM/RSEMTemp
module load gcc-4.9.2
module load python
module load jdk
module load perl
module load bowtie2

mkdir {srr}
cd {srr}
        '''.format(srr=_srr))
        
        if 'P' in _p:
            wget = srr2wgetEBI(_srr)
            fout.write('''
{wget} 
java -jar /panfs/panfs.cluster/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic/trimmomatic.jar  SE -threads 6 -phred33 {srr}_1.fastq.gz {srr}_1.fastq.trim ILLUMINACLIP:/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic-0.32/adapters/TruSeq3-PE.fa:2:30:10 SLIDINGWINDOW:4:20 LEADING:20 TRAILING:20 MINLEN:50 &&
rm {srr}_1.fastq.gz &
java -jar /panfs/panfs.cluster/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic/trimmomatic.jar  SE -threads 6 -phred33 {srr}_2.fastq.gz {srr}_2.fastq.trim ILLUMINACLIP:/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic-0.32/adapters/TruSeq3-PE.fa:2:30:10 SLIDINGWINDOW:4:20 LEADING:20 TRAILING:20 MINLEN:50 &&
rm {srr}_2.fastq.gz 
wait $(jobs -p)
/panfs/panfs.cluster/home/ks2073/p/2016/RSEM-1.3.0/rsem-calculate-expression -p 12 --bowtie2 --no-bam-output {srr}_1.fastq.trim,{srr}_2.fastq.trim /panfs/panfs.cluster/scratch/ks2073/RSEM/RSEMTemp/DNAseq/ProDB {srr}
            '''.format(srr = _srr, srr6 = _srr[:6], wget = wget))
        else:
            wget = srr2wgetEBI(_srr,False)
            fout.write('''
{wget} 
java -jar /panfs/panfs.cluster/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic/trimmomatic.jar  SE -threads 6 -phred33 {srr}.fastq.gz {srr}.fastq.trim ILLUMINACLIP:/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic-0.32/adapters/TruSeq3-PE.fa:2:30:10 SLIDINGWINDOW:4:20 LEADING:20 TRAILING:20 MINLEN:50 &&
rm {srr}.fastq.gz 
/panfs/panfs.cluster/home/ks2073/p/2016/RSEM-1.3.0/rsem-calculate-expression -p 12 --bowtie2 --no-bam-output {srr}.fastq.trim /panfs/panfs.cluster/scratch/ks2073/RSEM/RSEMTemp/DNAseq/ProDB {srr}
            '''.format(srr = _srr, srr6 = _srr[:6],wget = wget))
        fout.write('''\nrm *.fastq.trim\n''')
        fout.close()
#        break
    
    #download from ncbi
    lsSrrNCBI=[]
    def srr2wgetNCBI(srr,paired = True):
        '''
        get the download scripts from NCBI
        '''
        _srr = srr
        txt1 = '\nwget ftp://ftp-trace.ncbi.nih.gov/sra/sra-instant/reads/ByRun/sra/'+_srr[:3]+'/'+_srr[:6]+'/'+_srr+'/'+_srr+'.sra\n'
        txt2 = "/panfs/panfs.cluster/home/ks2073/p/2016/sratoolkit.2.8.0-ubuntu64/bin/fastq-dump --gzip --defline-seq '@$sn[_$rn]/$ri'  --split-files ./%s.sra  &&\nrm ./%s.sra \n"%(_srr,_srr)
        if not paired:
            txt2 = txt2 + 'mv {srr}_1.fastq.gz {srr}.fastq.gz\n'.format(srr=srr)
        return txt1 + txt2
    for line in ls[1:]:
        _es = line.split('\t')
        _srr = _es[1]
        lsSrrNCBI.append(_srr)
        _p = _es[3]
        fout = open('2017NCBI_DmRSEM_%s.txt'%(_srr),'w')
        fout.write('''
#!/bin/bash
#PBS -q batch
#PBS -N {srr}
#PBS -l nodes=1:ppn=12
#PBS -l walltime=120:00:00
#PBS -m abe -M atps@outlook.com
#PBS -j oe
cd /panfs/panfs.cluster/scratch/ks2073/RSEM/RSEMTemp
module load gcc-4.9.2
module load python
module load jdk
module load perl
module load bowtie2

mkdir {srr}
cd {srr}

        '''.format(srr=_srr))
        
        if 'P' in _p:
            fout.write(srr2wgetNCBI(_srr,True))
            fout.write('''

java -jar /panfs/panfs.cluster/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic/trimmomatic.jar  SE -threads 6 -phred33 {srr}_1.fastq.gz {srr}_1.fastq.trim ILLUMINACLIP:/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic-0.32/adapters/TruSeq3-PE.fa:2:30:10 SLIDINGWINDOW:4:20 LEADING:20 TRAILING:20 MINLEN:50 &&
rm {srr}_1.fastq.gz &
java -jar /panfs/panfs.cluster/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic/trimmomatic.jar  SE -threads 6 -phred33 {srr}_2.fastq.gz {srr}_2.fastq.trim ILLUMINACLIP:/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic-0.32/adapters/TruSeq3-PE.fa:2:30:10 SLIDINGWINDOW:4:20 LEADING:20 TRAILING:20 MINLEN:50 &&
rm {srr}_2.fastq.gz 
wait $(jobs -p)
/panfs/panfs.cluster/home/ks2073/p/2016/RSEM-1.3.0/rsem-calculate-expression -p 12 --bowtie2 --no-bam-output {srr}_1.fastq.trim,{srr}_2.fastq.trim /panfs/panfs.cluster/scratch/ks2073/RSEM/RSEMTemp/DNAseq/ProDB {srr}
            '''.format(srr = _srr, srr6 = _srr[:6]))
        else:
            fout.write(srr2wgetNCBI(_srr,False))
            fout.write('''

java -jar /panfs/panfs.cluster/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic/trimmomatic.jar  SE -threads 6 -phred33 {srr}.fastq.gz {srr}.fastq.trim ILLUMINACLIP:/home/ks2073/p/2016/trinityrnaseq-Trinity-v2.3.2/trinity-plugins/Trimmomatic-0.32/adapters/TruSeq3-PE.fa:2:30:10 SLIDINGWINDOW:4:20 LEADING:20 TRAILING:20 MINLEN:25 &&
rm {srr}.fastq.gz 
/panfs/panfs.cluster/home/ks2073/p/2016/RSEM-1.3.0/rsem-calculate-expression -p 12 --bowtie2 --no-bam-output {srr}.fastq.trim /panfs/panfs.cluster/scratch/ks2073/RSEM/RSEMTemp/DNAseq/ProDB {srr}
            '''.format(srr = _srr, srr6 = _srr[:6]))
        fout.write('''\nrm *.fastq.trim\n''')
        fout.close()
#        break

def dealwithDmTreeFile20170323():
    '''
    change add more information to gene name
    '''
    ls = '''FBpp0082539_392_S-Clip-SP FBpp0304587_528_S-Gd-SP FBpp0087783_271_S-SP FBpp0076693_2616_3LDLa-SP-3LDLa-SPH-3LDLa FBpp0082704_787_Clip-SP FBpp0084878_265_S-SP FBpp0312057_270_S-SP FBpp0084877_272_S-SP FBpp0082184_435_S-Clip-SP FBpp0087257_256_S-SP FBpp0087223_253_S-SP FBpp0087225_253_S-SP FBpp0087255_253_S-SP FBpp0087256_256_S-SP FBpp0087259_262_S-SP FBpp0087258_262_S-SP FBpp0087260_280_S-SP FBpp0073116_1047_S-5Clip-SPH FBpp0077483_245_S-SP FBpp0076964_259_S-SP FBpp0087226_252_S-SP FBpp0290879_267_SP FBpp0086785_260_S-SP FBpp0071322_397_S-Clip-SP FBpp0099851_266_S-SP FBpp0074969_273_S-SP FBpp0311236_2792_S-15CBD2-LDLa-SR-LDLa-SR-SP FBpp0070268_305_S-SP FBpp0070247_302_S-SP FBpp0082359_439_S-Gd-SP FBpp0078543_390_S-Clip-SP FBpp0289839_910_S-Clip-SPH-5repeats FBpp0080316_294_S-SPH FBpp0080315_464_S-Clip-SPH FBpp0080249_299_S-SPH FBpp0070826_308_S-SP FBpp0070828_362_S-SP FBpp0071118_1056_SEA-Frizzled-2LDLa-SPH FBpp0071175_393_S-Clip-SP FBpp0073978_276_S-SP FBpp0074138_251_S-SP FBpp0074137_249_S-SPH FBpp0074136_261_S-SPH FBpp0074095_254_S-SPH FBpp0074135_253_S-SPH FBpp0111913_637_S-Clip-SP FBpp0074400_394_S-Clip-SP FBpp0289579_286_S-SP FBpp0076963_260_S-SP FBpp0077758_271_S-SP FBpp0077757_277_S-SP FBpp0077526_270_S-SPH FBpp0077481_255_S-SP FBpp0077436_242_S-SPH FBpp0077373_343_S-Clip-SPH FBpp0289585_347_S-Clip-SPH FBpp0077068_314_S-SP FBpp0078692_258_S-SP FBpp0078691_276_S-SP FBpp0079653_406_S-Clip-SPH FBpp0080053_355_S-Clip-SPH FBpp0290909_299_S-SPH FBpp0080250_349_S-SP FBpp0080557_494_S-Clip-SPH FBpp0290124_379_S-SPH FBpp0080715_385_S-Clip-SPH FBpp0085397_655_S-3Clip-SPH FBpp0087983_1397_Frizzled-2LDLa-SR-SP FBpp0087871_529_S-CUB-SP FBpp0087804_444_S-2Clip-SPH FBpp0087784_459_S-2Clip-SP FBpp0271892_1674_Clip-SP FBpp0290411_561_S-Clip-SP FBpp0087747_389_S-Clip-SP FBpp0087746_855_Clip-SP FBpp0087551_317_S-SP FBpp0308511_332_S-SP FBpp0086427_260_S-SP FBpp0099998_362_S-Clip-SP FBpp0086075_523_S-SP-SPH FBpp0288561_279_S-SP FBpp0099932_267_S-SP FBpp0071657_372_S-SP FBpp0071636_352_S-SP FBpp0308966_292_S-SP FBpp0071815_360_S-Clip-SP FBpp0071898_265_S-SPH FBpp0072264_297_S-SPH FBpp0072287_249_S-SP FBpp0073064_322_S-SPH FBpp0073071_511_S-2Clip-SP FBpp0076776_272_S-SP FBpp0076747_319_S-SP FBpp0076749_274_S-SP FBpp0076750_259_S-SP FBpp0076775_262_S-SP FBpp0076752_438_S-SP FBpp0076774_290_S-SP FBpp0076772_267_S-SP FBpp0076540_258_S-SP FBpp0076335_260_S-SP FBpp0076389_262_S-SP FBpp0076159_315_S-SPH FBpp0076150_252_S-SPH FBpp0076116_259_S-SP FBpp0076117_268_S-SP FBpp0076118_264_S-SP FBpp0075737_287_S-SP FBpp0112080_748_S-2TSP-SP FBpp0305252_374_S-SP FBpp0075456_374_S-SP FBpp0075136_891_S-SPH FBpp0074946_270_S-SP FBpp0300968_285_S-SP FBpp0074838_322_S-SPH FBpp0301124_289_S-SPH FBpp0074713_408_S-Clip-SP FBpp0077991_275_S-SP FBpp0077990_292_S-SP FBpp0305481_289_S-SP FBpp0088629_392_S-SP FBpp0081237_391_S-Clip-SP FBpp0081424_405_S-SPH FBpp0081536_265_S-SP FBpp0081535_265_S-SP FBpp0291517_275_S-SP FBpp0082090_283_S-SP FBpp0082032_267_S-SP FBpp0082083_287_S-SP FBpp0082173_398_S-Clip-SP FBpp0289348_431_S-SP FBpp0082218_356_S-Clip-SP FBpp0082363_504_S-Gd-SP FBpp0082425_360_S-Clip-SP FBpp0110070_268_S-SP FBpp0082777_442_SP FBpp0082859_267_S-SP FBpp0082858_288_S-SP FBpp0292474_265_S-SP FBpp0082846_272_S-SP FBpp0082847_273_S-SP FBpp0083021_334_S-SP FBpp0290431_721_S-Clip-SP FBpp0292249_372_S-Clip-SP FBpp0083832_400_S-Clip-SP FBpp0311467_509_S-4Clip-SP FBpp0084143_333_S-SP FBpp0084482_377_S-Clip-SP FBpp0084484_381_S-Clip-SP FBpp0084629_265_S-SP FBpp0084731_319_S-SP FBpp0084732_316_S-SP FBpp0084897_253_S-SP FBpp0085001_424_S-Clip-SP FBpp0085000_418_S-Clip-SP FBpp0085052_267_S-SP FBpp0307429_367_S-Clip-SP FBpp0076746_270_S-SP FBpp0070158_254_S-SP FBpp0071658_364_S-SP FBpp0292250_341_S-Clip-SP FBpp0077482_248_S-SP FBpp0077433_251_S-SP FBpp0087222_272_S-SP FBpp0304273_263_S-SP FBpp0087254_253_S-SP FBpp0087224_253_S-SP FBpp0292043_280_S-SP FBpp0289610_279_S-SP FBpp0100028_277_S-SP FBpp0271804_281_S-SP FBpp0086417_291_S-SP FBpp0086415_526_S-SP-SPH FBpp0086312_264_S-SP FBpp0291139_530_S-SP-SPH FBpp0290173_277_S-SP FBpp0071612_284_S-SP FBpp0112279_282_S-SP FBpp0071614_316_S-SP FBpp0271829_306_S-SPH FBpp0087874_399_S-CUB-SP FBpp0087875_398_S-CUB-SP FBpp0291522_305_S-SP FBpp0083338_293_S-SPH FBpp0083337_284_S-SPH FBpp0289679_274_S-SPH FBpp0082798_628_S-4LDLa-2SUSHI-SP FBpp0112312_345_S-SP FBpp0291554_363_Clip-SP FBpp0082843_266_S-SP FBpp0289337_282_S-SPH FBpp0082845_275_S-SPH FBpp0082362_520_S-Gd-SP FBpp0077479_264_S-SP FBpp0080317_464_S-Clip-SPH FBpp0080318_294_S-SPH FBpp0077252_277_S-SP FBpp0073070_575_S-Clip-SP FBpp0100145_310_S-SP FBpp0099838_259_S-SP FBpp0072967_248_S-SP FBpp0072965_261_S-SP FBpp0076539_299_S-SP FBpp0076538_303_S-SP FBpp0289591_253_S-SPH FBpp0289592_253_S-SPH FBpp0076886_262_S-SP FBpp0070827_315_S-SP FBpp0070282_284_S-SP FBpp0071899_268_S-SPH FBpp0292307_504_S-SPH FBpp0077712_284_S-SPH FBpp0072966_257_S-SP FBpp0072964_258_S-SP FBpp0289634_293_S-SP FBpp0289635_516_S-Gd-SP FBpp0086233_281_S-SP FBpp0088798_284_S-SP FBpp0100103_275_S-SPH FBpp0291690_287_S-SP FBpp0086420_300_S-SP FBpp0290270_488_S-SPH-SPH FBpp0112536_421_S-Clip-SPH FBpp0300809_292_S-SP FBpp0110170_314_S-SPH FBpp0110171_297_S-SPH FBpp0293603_292_S-SPH FBpp0288806_282_S-SP FBpp0304083_511_S-SP FBpp0111649_270_S-SPH FBpp0111650_266_S-SPH FBpp0111676_257_S-SP FBpp0076748_271_S-SP FBpp0077480_245_S-SPH FBpp0080908_258_S-SP FBpp0291260_311_S-SP FBpp0290172_273_S-SP FBpp0292027_512_S-SPH-SPH FBpp0293436_483_S-SP-SPH FBpp0293647_245_S-SPH FBpp0293648_258_S-SPH FBpp0297782_281_S-SP FBpp0297783_279_S-SP FBpp0079916_270_S-SP FBpp0291994_265_S-SP FBpp0306131_474_S-SP-SPH FBpp0080214_239_S-SP FBpp0111472_1002_S-Clip-SP FBpp0294007_639_S-Clip-SP'''.split()
    txt = open(r"D:\Insects\Drosophila_melanogaster\SerineProteases\20170320Tree\20170321DmSPs249.nexus.con2.tre").read()
    for ele in ls:
        txt = txt.replace(ele.split('_')[0],ele)
    open(r"D:\Insects\Drosophila_melanogaster\SerineProteases\20170320Tree\20170321DmSPs249.nexus.con2.tre",'w').write(txt)
    
    txt = open(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\Fig\2017DmSPsChrLocation.svg").read()
    for ele in ls:
        txt = txt.replace(ele.split('_')[0],ele)
    open(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\Fig\2017DmSPsChrLocation.svg",'w').write(txt)


def RSEM_resultDealWith108Libraries20170412():
    '''
    extract data from RSEM runnings
    108 libraries
    '''
    #extract Trimmomatic survival rates and bowtie2 align rates
    folder = 'D:\\Insects\\Drosophila_melanogaster\\FPKM\\20170402RSEM_108libraries\\runningLog\\'
    import os
    files = os.listdir(folder)
    def getTrimmaticBowtie2Info(filename):
        '''
        given a filename, return Input reads, Suriving reads, and bowtie2 alignment rate
        like: [9676971,6816856,'93.92%']
        '''
        total = 0
        survive = 0
        rates = ''
        for line in open(filename):
            if line[:12] == 'Input Reads:':
                eles = line.split()
                total += int(eles[2])
                survive += int(eles[4])
            if 'overall alignment rate' in line:
                rates = line.split()[0]
        return [total, survive, rates]
    fout = open('list.txt','w')
    for file in files:
        info = getTrimmaticBowtie2Info(folder+file)
        info = [str(e) for e in info]
        fout.write(file.split('.')[0]+'\t'+'\t'.join(info)+'\n')
    fout.close()
    
    ls = '''SRR767609
    SRR767610
    SRR767611
    SRR767612
    SRR767607
    SRR100280
    SRR100279
    SRR100281
    SRR100282
    SRR100283
    '''.split()
    for file in files:
        srr = file.split('.')[0]
        if srr in ls:
            for line in open(folder+file):
                if ' reads; of these:' in line:
                    survive = line.split()[0]
                    print(srr,survive)
    
    
    #summarize RSEM running result
    import pandas as pd
    folderRSEM = 'D:\\Insects\\Drosophila_melanogaster\\2017RSEM_FPKM\\result\\'
    import os
    files = os.listdir(folderRSEM)
    f_libraryInfo = r"D:\Insects\Drosophila_melanogaster\2017RSEM_FPKM\libraryInfo.txt"
    df_lib = pd.read_csv(f_libraryInfo,sep = '\t',header = 0, index_col = 1)
    
    def getColRSEM(filename,column = 6):
        '''
        return column of a RSEM output file
        '''
        import pandas as pd
        df = pd.read_csv(filename,sep='\t')
        return df.iloc[:,column]
    dc = {}
    dc['transcripts'] = getColRSEM(folderRSEM+files[0],0)
    dc['length'] = getColRSEM(folderRSEM+files[0],2)
    for file in files:
        srr = file.split('.')[0]
        dc[srr] = getColRSEM(folderRSEM+file)
    df = pd.DataFrame(dc)
    df.head()
    #order
    index = ['transcripts','length'] +list(df_lib.index)
    df = df.loc[:,index]
    df.to_csv(folderRSEM+'..\\20170402RSEM_Dm6.14.csv')
    
    #from transcript to get gene FPKM values
    import pandas as pd
    df_xls = pd.read_excel(r"D:\Insects\Drosophila_melanogaster\2017RSEM_FPKM\20170402RSEM_Dm6.14.xlsx",index_col = 0)
    df_transcript = df_xls.copy()
    genes = df_xls.loc[:,'gene'].unique()[1:]
    df_genes = df_xls.iloc[:1,:].copy()
#    from collections import Counter
#    geneCount = Counter(list(df_xls.loc[:,'gene']))
    for n,gene in enumerate(genes):
        _dftemp = df_transcript[df_transcript.loc[:,'gene']==gene]
        df_genes.loc[n] = _dftemp.sum()
        df_genes.loc[n,'gene'] = gene
        df_genes.loc[n,'transcripts'] = ','.join(list(_dftemp.loc[:,'transcripts']))
        df_genes.loc[n,'length'] = 'na'
        print(df_genes.shape)
    df_genes.to_csv(r"D:\Insects\Drosophila_melanogaster\2017RSEM_FPKM\20170402RSEM_Dm6.14gene.csv")
    
    #the other way
    dftemp = df_xls.groupby('gene')
    ls_gene = []
    for name, group in dftemp:
#        print(name)
#        print(group)
        s = group.sum()
        s['gene'] = name
        s['transcripts'] = s['transcripts'].replace('FBtr',',FBtr')[1:]
        s['length'] = 'na'
        ls_gene.append(s)
    dfgenes = pd.concat(ls_gene,axis=1).transpose()
    dfgenes.to_csv(r"D:\Insects\Drosophila_melanogaster\2017RSEM_FPKM\20170402RSEM_Dm6.14gene.csv")
    
    

def DmSPSPH_FPKM_values_20170403():
    '''
    deal with DmSPSPHs. Use gene name to retrieve FPKM values
    '''
    import HeatmapCluster
    import pandas as pd
    df = pd.read_excel(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170403DmSPSPH_FPKM.xlsx",header = 1, index_col = 1)
    df = df.drop('transcripts',1)
    HeatmapCluster.fpkmHierarchicalCluster(df,'20170403DmSPSPHs_',col_cluster=False, row_cluster=True)

def DmSpnFindSeqs20170409():
    info = '''Spn27A CG11331 
    Spn28B CG6717 
    Spn28Da CG31902 L/S
    Spn28Db CG33121
    Spn28Dc CG7219 S/G
    Spn28F CG8137 Y/S
    Spn31A CG4804
    Spn38F CG9334 K/S
    Spn42DaA CG9453 R/A
    Spn42Db CG9454 K/G
    Spn42Dc CG9455 M/M
    Spn42Dd CG9456 R/A
    Spn42De CG9460 E/S
    Spn43Aa CG12172 M/S
    Spn43Ab CG1865
    Spn43Ac CG1857 L/S 
    Spn43Ad CG1859
    Spn47C CG7722
    Spn53F CG10956
    Spn55B CG10913 R/M
    Spn75F CG32203
    Spn76A CG3801
    Spn77Ba CG6680 K/A
    Spn77Bb CG6663
    Spn77Bc CG6289
    Spn85F CG12807
    Spn88Ea CG18525
    Spn88Eb CG6687
    Spn100A CG1342
    Spn? CG14470
    '''
    import re
    spnids = re.findall('CG\d*',info)
    import largeFastaFile
    ls = largeFastaFile.open_fasta_to_list(r"D:\Insects\Drosophila_melanogaster\dmel-all-translation-r6.14.fasta")
    fout = open('list.txt','w')
    for _s in ls:
        _cgs = re.findall('CG\d*',_s.description)
        _i = False
        for _cg in _cgs:
            if _cg in spnids:
                _i = True
                break
        if _i:
            fout.write('>'+_s.id+'_'+_cg+'\n'+str(_s.seq)+'\n')
    fout.close()
    
    
def checkmoreSPSPH20170530():
    '''
    serine proteases were checked from smart domain structure.
    some serine protease domain is modeled by pfam or Interpro only, not by smart
    chech to see whether more genes are missing in the original list
    '''
    ls_smart = open(r"D:\Insects\Drosophila_melanogaster\SerineProteases\list.txt").read().split()
    st_smart = set(ls_smart)
    import pandas as pd
    df = pd.read_csv(r"D:\Insects\Drosophila_melanogaster\interpro\DmeInterpro.tsv",sep = '\t',header = None, names = range(15))
    df_sp = df[df.iloc[:,12].apply(lambda x: 'Serine protease' in str(x) or 'serine protease' in str(x))]
    st_all = set(df_sp.iloc[:,0])
    st_new = st_all - st_smart
    
    #get sequences
    from Bio import SeqIO
    dcFa = SeqIO.to_dict(SeqIO.parse(open(r"D:\Insects\Drosophila_melanogaster\dmel-all-translation-r6.09.fasta"),"fasta"))
    fout = open(r"D:\Insects\Drosophila_melanogaster\SerineProteases\SPSPH_aaNew.txt","w")
    for ele in st_new:
        if ele in dcFa:
            fout.write(">"+dcFa[ele].id +" "+dcFa[ele].description.split()[4][5:-1]+" "+dcFa[ele].description.split()[5]\
            +" " +dcFa[ele].description.split()[8]+"\n"+str(dcFa[ele].seq)+"\n")
    fout.close()

def muscleAlginementCleanUp20170601():
    '''
    203 Clip proteases from Dm, Ms and Ag, multiple sequence alignment by muscle
    try to remove 'strange' sequences, which have less shared regions compared to others
    '''
    from Bio import SeqIO
    #read in alignments
    ls_alignments = list(SeqIO.parse(open(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170530Clips_Ms_Ag_Dm_Tree\2017ClipSeqs.fasta"),'fasta'))
    import pandas as pd
    def change(i):
        if str(i) =='-':
            return 0
        return 1
    df = pd.DataFrame([[change(i) for i in e.seq] for e in ls_alignments], index = [e.id for e in ls_alignments])
    #203 rows, 3546 columns
    
    
    max_missing = 1 #keep sites with no missing values
    residues_eachSite = [sum(df.iloc[:,i]) for i in range(3546)]
    site_keep = [n for n in range(3546) if residues_eachSite[n] >= 203-max_missing]
    df_keep = df.iloc[:,site_keep]
    seq_keep = []
    for _seqid in df_keep.index:
        if sum(df_keep.loc[_seqid,:]) == df_keep.shape[1]:
            seq_keep.append(_seqid)
    print(df_keep.shape[1],len(seq_keep))
    
    #save kept seqs
    fout = open(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170530Clips_Ms_Ag_Dm_Tree\2017ClipSeqs_keep192.fasta",'w')
    fout2 = open(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170530Clips_Ms_Ag_Dm_Tree\2017ClipSeqs_removed192.fasta",'w')
    for seq in ls_alignments:
        _towrite = '>'+seq.id+'\n'+str(seq.seq).replace('-','')+'\n'
        if seq.id in seq_keep:
            fout.write(_towrite)
        else:
            fout2.write(_towrite)
    fout.close()
    fout2.close()

def plot_FPKM_20170601():
    import HeatmapCluster
    import pandas as pd
    df = pd.read_csv(r"list.txt",sep = '\t', index_col = 0)
    print(df.shape)
    HeatmapCluster.fpkmHierarchicalCluster(df,'20170601DmSPSPH_',col_cluster=False, row_cluster=True)

def chromosomal_location20170605():
    st_allSPs = '''FBpp0082539 FBpp0304587 FBpp0087783 FBpp0076693 FBpp0082704 FBpp0084878 FBpp0312057 FBpp0084877 FBpp0082184 FBpp0087257 FBpp0087223 FBpp0087225 FBpp0087255 FBpp0087256 FBpp0087259 FBpp0087258 FBpp0087260 FBpp0073116 FBpp0077483 FBpp0076964 FBpp0087226 FBpp0290879 FBpp0086785 FBpp0071322 FBpp0099851 FBpp0074969 FBpp0311236 FBpp0070268 FBpp0070247 FBpp0082359 FBpp0078543 FBpp0289839 FBpp0080316 FBpp0080315 FBpp0080249 FBpp0070826 FBpp0070828 FBpp0071118 FBpp0071175 FBpp0073978 FBpp0074138 FBpp0074137 FBpp0074136 FBpp0074095 FBpp0074135 FBpp0074401 FBpp0074400 FBpp0289579 FBpp0076963 FBpp0077758 FBpp0077757 FBpp0077526 FBpp0077481 FBpp0077436 FBpp0077373 FBpp0289585 FBpp0077068 FBpp0078692 FBpp0078691 FBpp0079653 FBpp0080053 FBpp0290909 FBpp0080250 FBpp0080557 FBpp0290124 FBpp0080715 FBpp0085397 FBpp0087983 FBpp0087871 FBpp0087804 FBpp0087784 FBpp0271892 FBpp0290411 FBpp0087747 FBpp0087746 FBpp0087551 FBpp0308511 FBpp0086427 FBpp0099998 FBpp0086075 FBpp0288561 FBpp0099932 FBpp0071657 FBpp0071636 FBpp0308966 FBpp0071815 FBpp0071898 FBpp0072264 FBpp0072287 FBpp0073064 FBpp0073071 FBpp0076776 FBpp0076747 FBpp0076749 FBpp0076750 FBpp0076775 FBpp0076752 FBpp0076774 FBpp0076772 FBpp0076540 FBpp0076335 FBpp0076389 FBpp0076159 FBpp0076150 FBpp0076116 FBpp0076117 FBpp0076118 FBpp0075737 FBpp0112080 FBpp0305252 FBpp0075456 FBpp0075136 FBpp0074946 FBpp0300968 FBpp0074838 FBpp0301124 FBpp0074713 FBpp0077991 FBpp0077990 FBpp0305481 FBpp0088629 FBpp0081237 FBpp0081424 FBpp0081536 FBpp0081535 FBpp0291517 FBpp0082090 FBpp0082032 FBpp0082083 FBpp0082173 FBpp0289348 FBpp0082218 FBpp0082363 FBpp0082425 FBpp0110070 FBpp0082777 FBpp0082859 FBpp0082858 FBpp0292474 FBpp0082846 FBpp0082847 FBpp0083021 FBpp0290431 FBpp0292249 FBpp0083832 FBpp0311467 FBpp0084143 FBpp0084482 FBpp0084484 FBpp0084629 FBpp0084731 FBpp0084732 FBpp0084897 FBpp0085001 FBpp0085000 FBpp0085052 FBpp0085014 FBpp0307429 FBpp0076746 FBpp0070158 FBpp0071658 FBpp0292250 FBpp0077482 FBpp0077433 FBpp0087222 FBpp0304273 FBpp0087254 FBpp0087224 FBpp0292043 FBpp0289610 FBpp0100028 FBpp0271804 FBpp0086417 FBpp0086415 FBpp0086312 FBpp0291139 FBpp0290173 FBpp0071612 FBpp0112279 FBpp0071614 FBpp0271829 FBpp0087874 FBpp0087875 FBpp0291522 FBpp0083338 FBpp0083337 FBpp0289679 FBpp0082798 FBpp0112312 FBpp0291554 FBpp0082843 FBpp0289337 FBpp0082845 FBpp0082362 FBpp0077479 FBpp0080317 FBpp0080318 FBpp0077252 FBpp0073070 FBpp0100145 FBpp0099838 FBpp0072967 FBpp0072965 FBpp0076539 FBpp0076538 FBpp0289591 FBpp0289592 FBpp0076886 FBpp0070827 FBpp0070282 FBpp0071899 FBpp0292307 FBpp0077712 FBpp0072966 FBpp0072964 FBpp0289634 FBpp0289635 FBpp0086233 FBpp0088798 FBpp0100103 FBpp0291690 FBpp0086420 FBpp0290270 FBpp0112536 FBpp0300809 FBpp0110170 FBpp0110171 FBpp0293603 FBpp0288806 FBpp0304083 FBpp0111649 FBpp0111650 FBpp0111676 FBpp0076748 FBpp0077480 FBpp0080908 FBpp0291260 FBpp0290172 FBpp0292027 FBpp0293436 FBpp0293647 FBpp0293648 FBpp0297782 FBpp0297783 FBpp0079916 FBpp0291994 FBpp0306131 FBpp0080214 FBpp0111472 FBpp0294007 FBpp0084730 FBpp0112313 FBpp0084680 FBpp0084683 FBpp0111404 FBpp0070829 FBpp0084682'''.split()
    st_allSPs = set(st_allSPs)
    file_location = r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\Fig\2017DmSPsChrLocation.svg"
    txt_location = open(file_location,'r').read()
    import re
    st_SPfig = set(re.findall('FBpp\d{7}',txt_location))
    print(st_allSPs - st_SPfig)

def fpkm_fig20170606():
    filename = r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170601RSEM_expression\20170403DmSPSPH_FPKM_simplified.xlsx"
    import pandas as pd
    df = pd.read_excel(filename, skiprows = [0], index_col = 1)
    dfGP = df.loc[[ele for ele in df.index if df.loc[ele,'group'] == 'G'],:]
    dfGP = dfGP.iloc[:,6:]
    dfCLIP = df.loc[[ele for ele in df.index if df.loc[ele,'group'] == 'C'],:]
    dfCLIP = dfCLIP.iloc[:,6:]
    dfSPMP = df.loc[[ele for ele in df.index if df.loc[ele,'group'] == 'S' or df.loc[ele,'group'] == 'M' ],:]
    dfSPMP = dfSPMP.iloc[:,6:]
    
    import HeatmapCluster
    HeatmapCluster.fpkmHierarchicalCluster(dfGP,r'C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170601RSEM_expression\20170403DmGP_',col_cluster=False, row_cluster=True)
    HeatmapCluster.fpkmHierarchicalCluster(dfCLIP,r'C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170601RSEM_expression\20170403DmCLIP_',col_cluster=False, row_cluster=True)
    HeatmapCluster.fpkmHierarchicalCluster(dfSPMP,r'C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170601RSEM_expression\20170403DmSPMP_',col_cluster=False, row_cluster=True)

def changeChromosomelocationFig20170606():
    '''
    change FBpp number to annotation and protein SP name
    '''
    filename = r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\Fig\2017DmSPsChrLocationV2.svg"
    dc = {}
    for line in open(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\Fig\Fpp2Annotation.txt"):
        _k,_v = line.split('\t')
        _v = _v.replace('\n','')
        dc[_k] = _v
    txt = open(filename).read()
    for _k in dc:
        if _k in txt:
            txt = txt.replace(_k,dc[_k])
    open(filename,'w').write(txt)

def filterProteomeData20170607():
    '''
    keep lines with gene ids of SPs
    '''
    import pandas as pd
    df = pd.read_csv(r"D:\Insects\Drosophila_melanogaster\2017DmProteomeDevelopmental\proteinGroups20170417DeLFQ.csv")
    df2 = pd.read_excel(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170320DmSPSPHs.xlsx")
    genes = set(df2.loc[:,'gene'])
    ls = []
    for n in range(df.shape[0]):
        gs = df.iloc[n,0]
        if type(gs) is not str:
            continue
        gs = gs.split(';')
        for g in gs:
            if g in genes:
                ls.append(n)
                break
    df_sp = df.iloc[ls,:]
    df_sp.to_csv(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170607DmExpressionLFQ.csv")
    
    #plot
    dfLFQ = pd.read_excel(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170607DmExpressionLFQ.xlsx",index_col=0 )
    dfLFQ = dfLFQ.iloc[:,3:]
    import HeatmapCluster
    HeatmapCluster.fpkmHierarchicalCluster(dfLFQ,r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\20170607DmExpressionLFQ_",col_cluster=False, row_cluster=True)


def threewayAgreementsCount20170621():
    '''
    chr location, tree location expression groups
    '''
    import pandas as pd
    df = pd.read_csv('list.txt',sep='\t',header=None)
    GROUP = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ123456'
    for n in range(3):
        print('column',n,'in group number is', len([e for e in df.iloc[:,n] if e in GROUP]))
    dc = {}
    dc['CT'] =[]
    dc['CE'] = []
    dc['TE'] = []
    dc['CTE'] = []
    for n in range(257):
        info = ''.join(df.iloc[n,:])
        if info[0] in GROUP and info[1] in GROUP:
            dc['CT'].append(info[0]+info[1])
        if info[0] in GROUP and info[2] in GROUP:
            dc['CE'].append(info[0]+info[2])
        if info[1] in GROUP and info[2] in GROUP:
            dc['TE'].append(info[1]+info[2])
        if info[0] in GROUP and info[1] in GROUP and info[2] in GROUP:
            dc['CTE'].append(info)
    from collections import Counter
    dcCount = {}
    for key in dc:
        dcCount[key] = Counter(dc[key])
    for key in dcCount:
        print(key,sum([e for e in dcCount[key].values() if e>1]),len([e for e in dcCount[key].values() if e>1]))
    
def fourwayAgreementCount20170731():
    '''
    grouping, chr location, tree location expression groups
    '''
    import pandas as pd
    df = pd.read_csv(r"C:\Users\k\OneDrive\Lab\works\2017DmSPSPH\Fig\raw\20170731GCTE.txt",sep='\t')
    df.columns = list('NGCTE')
    GROUP = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890'
    for n in 'CTE':
        print('column',n,'in group number is', len([e for e in df.loc[:,n] if e in GROUP]))
    
    from collections import Counter
    print(Counter(df.loc[:,'G']))
    def CTEagreement(df):
        dc = {}
        dc['CT'] =[]
        dc['CE'] = []
        dc['TE'] = []
        dc['CTE'] = []
        for n in range(df.shape[0]):
            info = ''.join(df.iloc[n,2:])
            if info[0] in GROUP and info[1] in GROUP:
                dc['CT'].append(info[0]+info[1])
            if info[0] in GROUP and info[2] in GROUP:
                dc['CE'].append(info[0]+info[2])
            if info[1] in GROUP and info[2] in GROUP:
                dc['TE'].append(info[1]+info[2])
            if info[0] in GROUP and info[1] in GROUP and info[2] in GROUP:
                dc['CTE'].append(info)
        dcCount = {}
        for key in dc:
            dcCount[key] = Counter(dc[key])
        for key in dcCount:
            print(key,sum([e for e in dcCount[key].values() if e>1]),len([e for e in dcCount[key].values() if e>1]))
    CTEagreement(df)
    CTEagreement(df[df.loc[:,'G'] =='M'])
    CTEagreement(df[df.loc[:,'G'] =='S'])
    CTEagreement(df[df.loc[:,'G'] =='G'])
    CTEagreement(df[df.loc[:,'G'] =='C'])
    
    def CTE_in_GROUP_asExpected(row, expected = [1,1,1]):
        '''
        return True or False, to see if CTE in GROUP as expected
        '''
        if (row['C'] in GROUP) == expected[0] and (row['T'] in GROUP) == expected[1] and (row['E'] in GROUP) == expected[2]:
            return True
        return False

    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[1,1,1]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[1,1,0]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[1,0,1]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[0,1,1]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[1,0,0]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[0,1,0]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[0,0,1]), axis=1)].shape)
    print(df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,[0,0,0]), axis=1)].shape)
    
    from itertools import product
    ls_group = list(product([0,1],repeat = 3))
    for _e in ls_group:
        _df = df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,_e), axis=1)]
        print(_e, _df.shape)
        for _ee in "CMSG":
            print(_ee, _df[_df.loc[:,'G']==_ee].shape)
        
    for _e in ls_group:
        _df = df[df.apply(lambda x: CTE_in_GROUP_asExpected(x,_e), axis=1)]
        
    
    for _ee in 'CMSG':
        _df =  df[df.loc[:,'G']==_ee]
        print(_ee, _df.shape)
        for _e in ls_group:
            _dfdf = _df[_df.apply(lambda x: CTE_in_GROUP_asExpected(x,_e), axis=1)]
            print(_e,_dfdf.shape)
    
    #count CTE relationship
    #CT*, with CT number over2
    dc = {}
    dcGene = {}
    for check in ['CT','CE','TE','CTE']:
        dc[check] = df[df.apply(lambda x:all(x[_i] in GROUP for _i in check),axis = 1)]
        dcGene[check] = {}
        for _n in range(dc[check].shape[0]):
            _row = dc[check].iloc[_n,:]
            _key =''.join( _row[_i] for _i in check)
            if _key not in dcGene[check]:
                dcGene[check][_key] = []
            dcGene[check][_key].append((_row['N'], _row['G']))
        
        _ls = []
        for _e in dcGene[check].values():
            if len(_e) >1:
                _ls += _e
        print(check, len(_ls))
        for _i in 'CMSG':
            print(_i, len([e for e in _ls if e[1] == _i]))
    
    # save three way agreement
    fout = open('list.txt','w')
    for _k in dcGene['CTE']:
        if len(dcGene['CTE'][_k])>1:
            fout.write(_k + '\t'+','.join(_e[0]+' '+_e[1] for _e in dcGene['CTE'][_k]) +'\n')
    fout.close()
    
    def threeWayCountIndividual(df):
        dc = {}
        dcGene = {}
        for check in ['CT','CE','TE','CTE']:
            dc[check] = df[df.apply(lambda x:all(x[_i] in GROUP for _i in check),axis = 1)]
            dcGene[check] = {}
            for _n in range(dc[check].shape[0]):
                _row = dc[check].iloc[_n,:]
                _key =''.join( _row[_i] for _i in check)
                if _key not in dcGene[check]:
                    dcGene[check][_key] = []
                dcGene[check][_key].append((_row['N'], _row['G']))
            
            _ls = []
            for _e in dcGene[check].values():
                if len(_e) >1:
                    _ls += _e
            print(check, len(_ls))
        for i in dcGene['CTE'].values():
            if len(i)>1:
                print(','.join(e[0] for e in i))
    
    threeWayCountIndividual(df)
    for _k in 'CMSG':
        print(_k)
        threeWayCountIndividual(df[df.loc[:,'G'] == _k])
    
    