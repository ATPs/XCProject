# -*- coding: utf-8 -*-
"""
Created on Mon Jun  6 14:41:36 2016

@author: k
"""


def srpLike():
    from seqExtract import seqRegExtract
    file_arthopoda = r'F:\Insects\uniprotkb_arthropoda20160510'
    file_eukaryota = r'F:\Insects\uniprotkb_eukaryota20160510'
    import re
    re_str = re.compile('[RK][ABD-Z]{1,15}C[ABD-Z]{3}[ABD-Z]{4,6}G[ABD-Z]{1,2}C[ABD-Z]{1,15}$')
    seqRegExtract(file_arthopoda,re_str,'arthopodaSRP.txt','fasta',50,250)
    seqRegExtract(file_eukaryota,re_str,'eukaryotaSRP.txt','fasta',50,250)
    
def findSRPlikeInlepidoptera20160816():
    """
    eight species with protein models available, find all SRP like proteins.
    """
    from seqExtract import seqRegExtract
    import re
    re_str = re.compile('[RK][ABD-Z]{1,15}C[ABD-Z]{3}[ABD-Z]{4,6}G[ABD-Z]{1,2}C[ABD-Z]{1,15}$')
    folder = r'F:\Insects\lepidoptera\proteinFasta'
    import glob
    files = glob.glob(folder+'\\*')
    for file in files:
        seqRegExtract(file, re_str, file.split("\\")[-1],'fasta', 50, 250)
    
    
    
def findSRPlikeInlepidoptera20160904():
    from seqExtract import seqRegExtract
    import re
    re_str = re.compile('[RK][ABD-Z]{1,15}C[ABD-Z]{3}[ABD-Z]{4,6}G[ABD-Z]{1,2}C[ABD-Z]{1,15}$')
    file = r"F:\Insects\uniprotkb_arthropoda20160715.fasta"
    seqRegExtract(file, re_str, file.split("\\")[-1],'fasta', 50, 300)
    
    folder = 'D:\\mine\\OneDrive\\Lab\\works\\20150403ManducaInsectCytokinesPP_SRP_uENF\\20160904findSRPCytokinesAthropoda\\'
    file_all = folder + '20160904AllPossibleSRPs.fasta'
    file_signalP = folder + '20160904sigalPReport.txt'
#    from Bio import SeqIO
    #pick up those with signal peptides.
    import largeFastaFile
    ls_all = largeFastaFile.open_fasta_to_list(file_all)
    ls_signalP = []
    for line in open(file_signalP):
        ls_signalP.append(line.split())
    
    ls_sig = []
    for num in range(len(ls_all)):
        if ls_signalP[num][0] == ls_all[num].id.replace(":","_"):
            if ls_signalP[num][-3] == 'Y':
                ls_sig.append(ls_all[num])
        else:
            print(ls_all[num].id)
            break
    
    largeFastaFile.saveFastaListToFile(ls_sig, folder+'20160904SRP_like_withSignalP.fasta')
    
    ls_cystein2 = [] #for seq in ls_sig, save those with only 2 cysteins in non-signal peptide region
    ls_cysteinMore = []
    dc_signal = {}
    for ele in ls_signalP:
        dc_signal[ele[0].replace('_',":",1)] = int(ele[2])
    for ele in ls_sig:
        if str(ele.seq)[dc_signal[ele.id]-1:].count('C') == 2:
            ls_cystein2.append(ele)
        else:
            ls_cysteinMore.append(ele)
    
    dc_active = {}
    for ele in ls_sig:
        seq = str(ele.seq)
        start = re.search(re_str, seq).start()
        dc_active[ele.id] = start
    
    #save ls_cystein2 and ls_cysteinMore to a word file for manually check.
    largeFastaFile.saveFastaListToFile(ls_cystein2, folder+'20160904SRP_like_withSignalP_cystein2.fasta')
    largeFastaFile.saveFastaListToFile(ls_cysteinMore, folder+'20160904SRP_like_withSignalP_cysteinMorethan2.fasta')
    
    from docx import Document
    from docx.shared import RGBColor
    document = Document()
    for ele in ls_sig:
        if ele in ls_cystein2:
            p = document.add_paragraph()#fasta head
            run = p.add_run("\n\n\n>"+ele.description +' len:'+ str(len(ele.seq)))
            font = run.font
            font.color.rgb = RGBColor(0xFF,0x00,0x00)
            font.name = "Courier New"
        else:
            p = document.add_paragraph()#fasta head
            run = p.add_run("\n\n\n>"+ele.description +' len:' + str(len(ele.seq)))
            font = run.font
            font.color.rgb = RGBColor(0x00,0x00,0x00)
            font.name = "Courier New"
        
        run = p.add_run('\n')
        signalp = str(ele.seq)[:dc_signal[ele.id]]
        run = p.add_run(signalp)
        font = run.font
        font.color.rgb = RGBColor(0x00,0x99,0x00)
        font.name = "Courier New"
        middle = str(ele.seq)[dc_signal[ele.id]:dc_active[ele.id]+1]
        run = p.add_run(middle)
        font = run.font
        font.color.rgb = RGBColor(0x00,0x00,0x00)
        font.name = "Courier New"
        active = str(ele.seq)[dc_active[ele.id]+1:]
        run = p.add_run(active)
        font = run.font
        font.color.rgb = RGBColor(0x00,0x00,0xFF)
        font.name = "Courier New"
    
    document.save(folder+'20160904SRP_like.docx')
    
    dc_good = largeFastaFile.open_fasta_to_dic(folder + '20160905goodSRPlikes.txt')
    
    fout = open(folder + '20160905SRP_good_pep_1.txt','w')
    for ele in ls_sig:
        if ele.id in dc_good:
            fout.write('>'+ele.id+'\n'+str(ele.seq)[dc_active[ele.id]+1:] +'\n')
    fout.close()
    
    ls_good = largeFastaFile.open_fasta_to_list(r"D:\mine\OneDrive\Lab\works\20150403ManducaInsectCytokinesPP_SRP_uENF\20160904findSRPCytokinesAthropoda\20160905SRP_all_possibleCombined1.txt")
    ls_good_unique = largeFastaFile.fasta_within_seq_big(ls_good)
    largeFastaFile.saveFastaListToFile(ls_good_unique, folder + '20160905SRP_all_possibleCombined_159seq.txt')
    
    #use id to save the sequences
    fout = open(folder + '20160905SRP_all_possibleCombined_159seqID.txt','w')
    for num in range(len(ls_good_unique)):
        ele = ls_good_unique[num]
        fout.write('>%03d'%(num+1)+'\n'+str(ele.seq)+'\n')
    fout.close()
    
    #order seq in ls_good_unique according to the tree.
    ls_good_uniqueAlign = largeFastaFile.open_fasta_to_list(r"D:\mine\OneDrive\Lab\works\20150403ManducaInsectCytokinesPP_SRP_uENF\20160904findSRPCytokinesAthropoda\20160905SRP_all_possibleCombined_159seqIDAlign.txt")
    ls_order = open('list.txt').read().split()
    ls_order = [int(num) - 1 for num in ls_order]
    fout = open('ordered.txt','w')
    for num in ls_order:
        fout.write(str(ls_good_uniqueAlign[num].seq)+'\n')
    fout.close()
    
    for num in ls_order:
        print(ls_good_unique[num].id)


def findSRPlikeInHuman20161101():
    from seqExtract import seqRegExtract
    import re
    re_str = re.compile('[RK][ABD-Z]{1,15}C[ABD-Z]{3}[ABD-Z]{4,6}G[ABD-Z]{1,2}C[ABD-Z]{1,15}$')
    file = r"F:\Insects\HomoSapiens\protein.fa"
    seqRegExtract(file, re_str, file.split("\\")[-1],'fasta', 50, 300)
    file = r"F:\Insects\HomoSapiens\Gnomon_prot.fsa"
    seqRegExtract(file, re_str, file.split("\\")[-1],'fasta', 50, 300)
    
    #sequences stored in "F:\Insects\HomoSapiens\20161101SRPcandidates.txt"
    #keep unique sequences
    filename = r"F:\Insects\HomoSapiens\20161101SRPcandidates.txt"
    import largeFastaFile
    lsSRPs = largeFastaFile.open_fasta_to_list(filename)
    lsUniSRPs = largeFastaFile.fasta_uni_keepone(lsSRPs)
    largeFastaFile.saveFastaListToFile(lsUniSRPs,filename +'Unique')
    
    ls_all = list(lsUniSRPs)
    file_signalP = r"F:\Insects\HomoSapiens\20161101SignalP_report.txt"
    ls_signalP = []
    for line in open(file_signalP):
        ls_signalP.append(line.split())
    
    ls_sig = []
    for num in range(len(ls_all)):
        if ls_signalP[num][0] == ls_all[num].id.replace("|","_"):
            if ls_signalP[num][-3] == 'Y':
                ls_sig.append(ls_all[num])
        else:
            print(ls_all[num].id)
            break
    
    #save the file
    fo = open(filename +'Unique.Signal','w')
    for ele in ls_sig:
        fo.write(">"+ele.description+'\n'+str(ele.seq)+'\n')
    fo.close()
    
    
