# -*- coding: utf-8 -*-
"""
Created on Fri Jun  9 10:38:33 2017

@author: k
"""

def getSPseqsFromMCOTOGS2_20170609():
    import largeFastaFile
    ls = open(r"D:\Insects\ManducaSexta\Interpro\SPids.txt").read().split()
    dc = largeFastaFile.open_fasta_to_dic(r"D:\Insects\ManducaSexta\Interpro\ms_ogs_proteins.fa")
    dc.update(largeFastaFile.open_fasta_to_dic(r"D:\Insects\ManducaSexta\Interpro\Ms_MCOTprotein.txt"))
    lsSP = [dc[e] for e in ls]
    lsSPu = largeFastaFile.fasta_within_seq_big_withError(lsSP,0)
    lsSPu2 = largeFastaFile.fasta_within_seq_big_withError(lsSP)
    largeFastaFile.saveFastaListToFile(lsSPu2,r"D:\Insects\ManducaSexta\Interpro\MsSP.txt")
    #store the tsv domain
    idSPu2 = set([e.id for e in lsSPu2])
    f_SPdomain = open(r"D:\Insects\ManducaSexta\Interpro\MsSPdomain.txt",'w')
    lsDomain = open(r"D:\Insects\ManducaSexta\Interpro\Ms_MCOTprotein.tsv").readlines()+open(r"D:\Insects\ManducaSexta\Interpro\ms_ogs_proteins.tsv").readlines()
    for _line in lsDomain:
        if _line.split()[0] in idSPu2:
            f_SPdomain.write(_line)
    f_SPdomain.close()
    
    #copy domain files to new folder
    from shutil import copyfile
    for _id in idSPu2:
        if 'MCOT' in _id:
            copyfile("D:\\Insects\\ManducaSexta\\Interpro\\html_MCOT\\"+_id+'.html', "D:\\Insects\\ManducaSexta\\Interpro\\SP\\"+_id+'.html')
        else:
            copyfile("D:\\Insects\\ManducaSexta\\Interpro\\html_OGS2\\"+_id+'.html', "D:\\Insects\\ManducaSexta\\Interpro\\SP\\"+_id+'.html')
    
    
def save2DocxWithSmartDomain20170609():
    """
    save the text file of fasta format to word, with smart domains labeled.
    """
    
    fileSPfa = r"D:\Insects\ManducaSexta\Interpro\20170612NewSP_seqs.txt"
    filelsDomainSP = r"D:\Insects\ManducaSexta\Interpro\MsSPdomain.txt"
    fileSPorder = r"D:\Insects\ManducaSexta\Interpro\MsSPorder.txt"
    from Bio import SeqIO
    dcFa = SeqIO.to_dict(SeqIO.parse(open(fileSPfa),"fasta"))
    lsOrder = open(fileSPorder).read().split()
    lsFa = [dcFa[e] for e in lsOrder]
    lsdomain = open(filelsDomainSP).readlines()
    dcDomain ={}
    for ele in lsdomain:#store domain information in a dictionary
        ele_detail = ele.split("\t")
        if ele_detail[0] not in dcDomain:
            dcDomain[ele_detail[0]] =[]
        dcDomain[ele_detail[0]].append(ele_detail)
    dcSignalP = {}#store signalP location
    for key in dcDomain:
        for ele in dcDomain[key]:
            if "SignalP_EUK" == ele[3]:
                dcSignalP[ele[0]] = ["SignalP",int(ele[6])-1,int(ele[7]),0]
    dcSmart = {}#store SMART domain location
    for key in dcDomain:
        for ele in dcDomain[key]:
            if "SMART" == ele[3]:
                if key not in dcSmart:
                    dcSmart[key] = []
                dcSmart[key].append([ele[5],int(ele[6])-1,int(ele[7]),float(ele[8])])
    dcSlices ={}
    for key in dcFa:#combin dcSignalP and dcSmart
        dcSlices[key] = []
        if key in dcSignalP:
            dcSlices[key].append(dcSignalP[key])
        if key in dcSmart:
            for ele in dcSmart[key]:
                dcSlices[key].append(ele)
    for key in dcSlices:#sort by starting site
        dcSlices[key] = sorted(dcSlices[key],key = lambda x: x[1])
    
    dcSP = {}#store SP domain
    for key in dcDomain:
        for ele in dcDomain[key]:
            if "SM00020" == ele[4]:
                dcSP[key] = [ele[5],int(ele[6])-1,int(ele[7]),float(ele[8])]
    
    dcDocSlices ={}
    for key in lsFa:
        key = key.id
        dcDocSlices[key] = []
        if key in dcSignalP:
            if key in dcSP:
                if dcSignalP[key][2]>dcSP[key][1]:
                    dcSP[key][1] =dcSignalP[key][2]
                dcDocSlices[key] = [dcSignalP[key],["",dcSignalP[key][2], dcSP[key][1]],dcSP[key],["",dcSP[key][2],len(dcFa[key])]]
            else:
                dcDocSlices[key] = [dcSignalP[key],["",dcSignalP[key][2],len(dcFa[key])]]
        else:
            if key in dcSP:
                dcDocSlices[key] = [["",0, dcSP[key][1]],dcSP[key],["",dcSP[key][2],len(dcFa[key])]]
            else:
                dcDocSlices[key] = [["",0,len(dcFa[key])]]
    
                    
    from docx import Document
    from docx.shared import RGBColor
    document = Document()
    for key in lsFa:
        key = key.id
        p = document.add_paragraph()#fasta head
        run = p.add_run("\n\n\n>"+dcFa[key].description)
        font = run.font
        font.color.rgb = RGBColor(0x00,0x00,0x00)
        font.name = "Courier New"
        
        p = document.add_paragraph()#domains and signalP basic
        for ele in dcSlices[key]:
            run = p.add_run(ele[0]+"\t")
            font = run.font
            font.color.rgb = RGBColor(0x00,0x00,0x00)
            font.name = "Courier New"
        
        seq = str(dcFa[key].seq)
        
        p = document.add_paragraph()#domain detail
        for ele in dcSlices[key]:
            if ele not in dcDocSlices[key]:
                run = p.add_run(ele[0] +"\t"+seq[ele[1]:ele[2]]+"\n")
                font = run.font
                font.color.rgb = RGBColor(0x00,0x00,0xFF)
                font.name = "Courier New"
        
        p = document.add_paragraph()#write seq
        for ele in dcDocSlices[key]:
            if "SignalP" in ele[0]:
                run = p.add_run(seq[ele[1]:ele[2]])
                font = run.font
                font.color.rgb = RGBColor(0x00,0xFF,0x00)
                font.name = "Courier New"
            elif "protease" in ele[0]:
                run = p.add_run(seq[ele[1]:ele[2]])
                font = run.font
                font.color.rgb = RGBColor(0xFF,0x00,0x00)
                font.name = "Courier New"
            else:
                run = p.add_run(seq[ele[1]:ele[2]])
                font = run.font
                font.color.rgb = RGBColor(0x00,0x00,0x00)
                font.name = "Courier New"
#        break
    document.save("domainSP.docx")

def getSPidsFromAnnotatedDOCX20170609():
    txt = open('list.txt').readlines()
    lsSP = []
    for _l in txt:
        if len(_l)>1:
            if _l[0] == '>':
                lsSP.append(_l.split()[0][1:])
    print(len(lsSP))

    from shutil import copyfile
    for _id in lsSP:
        copyfile("D:\\Insects\\ManducaSexta\\Interpro\\SP\\"+_id+'.html', "D:\\Insects\\ManducaSexta\\Interpro\\SP257\\"+_id+'.html')

def compareCurrentSPSPHwithPrevisouSPSPHpublished20170612():
    '''
    got 257 SPSPHs. Previous found, 193
    compare and remove duplicated ones
    '''
    import largeFastaFile
    ls1 = largeFastaFile.open_fasta_to_list(r"D:\Insects\ManducaSexta\Interpro\20160612MsSPSPH257MCOTOGS2.txt")
    ls2 = largeFastaFile.open_fasta_to_list(r"D:\Insects\ManducaSexta\Interpro\20170612_MsSPSPH193Annotated.txt")
    ls = ls1+ls2
    lsU = largeFastaFile.fasta_within_seq_big_withError(ls)
    lsU2 = largeFastaFile.fasta_uni_keepone(ls)
    
    lsNew = open(r"D:\Insects\ManducaSexta\Interpro\20170612NewSPids.txt").read().split()
    lsNewseq = [e for e in ls1 if e.id in lsNew]
    largeFastaFile.saveFastaListToFile(lsNewseq,r"D:\Insects\ManducaSexta\Interpro\20170612NewSP_seqs.txt",'tab')
    fout = open(r"D:\Insects\ManducaSexta\Interpro\20170612NewSP_seqs.txt",'w')
    
    from shutil import copyfile
    for _id in lsNew:
        if 'MCOT' in _id:
            copyfile("D:\\Insects\\ManducaSexta\\Interpro\\html_MCOT\\"+_id+'.html', "D:\\Insects\\ManducaSexta\\Interpro\\SP57\\"+_id+'.html')
        else:
            copyfile("D:\\Insects\\ManducaSexta\\Interpro\\html_OGS2\\"+_id+'.html', "D:\\Insects\\ManducaSexta\\Interpro\\SP57\\"+_id+'.html')